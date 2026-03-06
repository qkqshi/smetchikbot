
import docx
import os

def inspect_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    doc = docx.Document(file_path)
    print(f"Total tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}:")
        for row_idx, row in enumerate(table.rows[:3]):  # Look at first 3 rows
            cells_text = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            print(f"  Row {row_idx}: {cells_text}")

if __name__ == "__main__":
    inspect_docx(r"c:\Users\danil\Desktop\ворк\разработка\smetchikbot\uploads\ВЯЗНИКОВСКАЯ просчет.docx")

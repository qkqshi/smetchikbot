import docx
import re
import os
from docx.oxml import parse_xml

def extract_image_from_cell(cell, document, output_folder="uploads/images"):
    """
    Извлекает изображение из ячейки таблицы
    
    Args:
        cell: ячейка таблицы docx
        document: объект документа docx
        output_folder: папка для сохранения изображений
    
    Returns:
        путь к сохраненному изображению или None
    """
    os.makedirs(output_folder, exist_ok=True)
    
    # Ищем изображения в параграфах ячейки
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # Проверяем наличие изображений в run
            for drawing in run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                # Ищем blip (ссылку на изображение)
                blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                for blip in blips:
                    # Получаем ID изображения
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                        try:
                            # Получаем изображение из документа через document.part
                            image_part = document.part.related_parts[embed]
                            image_bytes = image_part.blob
                            
                            # Определяем расширение файла по content_type
                            content_type = image_part.content_type
                            ext = 'png'  # По умолчанию PNG
                            if 'jpeg' in content_type or 'jpg' in content_type:
                                ext = 'jpg'
                            elif 'png' in content_type:
                                ext = 'png'
                            elif 'gif' in content_type:
                                ext = 'gif'
                            elif 'bmp' in content_type:
                                ext = 'bmp'
                            
                            # Сохраняем изображение с правильным расширением
                            image_filename = f"{output_folder}/image_{embed.replace(':', '_')}.{ext}"
                            
                            # Используем Pillow для конвертации и валидации
                            try:
                                from PIL import Image
                                import io
                                
                                # Открываем изображение из байтов
                                img = Image.open(io.BytesIO(image_bytes))
                                
                                # Конвертируем в RGB если нужно (для JPEG)
                                if ext == 'jpg' and img.mode in ('RGBA', 'LA', 'P'):
                                    # Создаем белый фон
                                    background = Image.new('RGB', img.size, (255, 255, 255))
                                    if img.mode == 'P':
                                        img = img.convert('RGBA')
                                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                                    img = background
                                
                                # Сохраняем изображение
                                img.save(image_filename, format=ext.upper() if ext != 'jpg' else 'JPEG')
                                
                                return image_filename
                            except Exception as pil_error:
                                # Если Pillow не смог обработать, сохраняем как есть
                                print(f"⚠️ Pillow не смог обработать изображение: {pil_error}, сохраняю как есть")
                                with open(image_filename, 'wb') as f:
                                    f.write(image_bytes)
                                return image_filename
                        except Exception as e:
                            print(f"Ошибка извлечения изображения: {e}")
    return None

def extract_furniture_data(file_path):
    """
    Извлекает данные из единой таблицы (материалы и мебель вместе)
    """
    # ПРОВЕРКА 1: Существует ли файл вообще?
    if not os.path.exists(file_path):
        print(f"❌ ОШИБКА: Файл не найден по пути: {os.path.abspath(file_path)}")
        return [], []

    print(f"✅ Файл найден, начинаю чтение: {file_path}")
    doc = docx.Document(file_path)
    extracted_items = []
    furniture_items = []

    # ПРОВЕРКА 2: Есть ли в документе таблицы?
    if not doc.tables:
        print("⚠️ ПРЕДУПРЕЖДЕНИЕ: В документе не найдено ни одной таблицы!")
        return [], []

    print(f"🔍 Найдено таблиц в файле: {len(doc.tables)}")

    # ЕДИНАЯ ТАБЛИЦА: Материалы и мебель
    table = doc.tables[0]
    print(f"--- Обработка таблицы (Материалы и Мебель) ---")
    
    # Определяем индексы столбцов из заголовка
    header_row = table.rows[0]
    column_indices = {}
    for idx, cell in enumerate(header_row.cells):
        header_text = cell.text.strip().lower()
        if 'визуал' in header_text or 'фото' in header_text or 'изображ' in header_text:
            column_indices['visual'] = idx
        elif 'изделие' in header_text or 'название' in header_text or 'наименование' in header_text:
            if 'name' not in column_indices:  # Берем первое вхождение
                column_indices['name'] = idx
        elif 'количество' in header_text or 'кол-во' in header_text or 'кол' in header_text:
            column_indices['quantity'] = idx
        elif 'размер' in header_text or 'габарит' in header_text:
            if 'size' not in column_indices:  # Берем первое вхождение
                column_indices['size'] = idx
        elif 'корпус' in header_text:
            if 'body' not in column_indices:  # Берем первое вхождение
                column_indices['body'] = idx
        elif 'фасад' in header_text:
            if 'facade' not in column_indices:  # Берем первое вхождение
                column_indices['facade'] = idx
    
    print(f"📋 Найденные индексы столбцов: {column_indices}")
    
    for row_index, row in enumerate(table.rows):
        # Пропускаем заголовок
        if row_index == 0:
            continue
            
        cells = row.cells
        
        # Используем найденные индексы или значения по умолчанию
        name_idx = column_indices.get('name', 0)
        visual_idx = column_indices.get('visual', 2)
        quantity_idx = column_indices.get('quantity', 3)
        size_idx = column_indices.get('size', 4)
        body_idx = column_indices.get('body', 6)
        facade_idx = column_indices.get('facade', 10)
        
        item_name = cells[name_idx].text.strip()
        
        # Если первая ячейка пустая, пропускаем строку
        if not item_name:
            continue

        print(f"👉 Нашел изделие: {item_name}")
        
        # Извлекаем изображение из столбца "визуал"
        image_path = None
        if visual_idx < len(cells):
            image_path = extract_image_from_cell(cells[visual_idx], doc)
            if image_path:
                print(f"   🖼️ Извлечено изображение: {image_path}")
        
        # Извлекаем количество
        quantity = 1
        if quantity_idx < len(cells):
            quantity_text = cells[quantity_idx].text.strip()
            if quantity_text and quantity_text.isdigit():
                quantity = int(quantity_text)
                print(f"   📊 Количество: {quantity}")
        
        # Извлекаем размеры - проверяем несколько столбцов
        size_text = ""
        for possible_size_idx in [size_idx, size_idx + 1, size_idx + 2]:
            if possible_size_idx < len(cells):
                temp_text = cells[possible_size_idx].text.strip()
                # Проверяем, содержит ли текст размеры (ширина/высота/глубина)
                if any(keyword in temp_text.lower() for keyword in ['ширина', 'высота', 'глубина', 'width', 'height', 'depth']):
                    size_text = temp_text.lower()
                    print(f"   📏 Размеры найдены в столбце {possible_size_idx}")
                    break
        
        if not size_text:
            size_text = cells[size_idx].text.lower() if size_idx < len(cells) else ""
        
        print(f"   📏 Исходный текст размеров: {repr(size_text)}")
        
        width = 0
        height = 0
        depth = 600  # Значение по умолчанию
        
        try:
            # Пытаемся найти размеры по ключевым словам
            width_match = re.search(r'(?:ширина|width)[:\s]*(\d+)', size_text, re.IGNORECASE)
            height_match = re.search(r'(?:высота|height)[:\s]*(\d+)', size_text, re.IGNORECASE)
            depth_match = re.search(r'(?:глубина|depth)[:\s]*(\d+)', size_text, re.IGNORECASE)
            
            if width_match and height_match:
                # Найдены размеры по ключевым словам
                width = int(width_match.group(1))
                height = int(height_match.group(1))
                if depth_match:
                    depth = int(depth_match.group(1))
                print(f"   ✅ Размеры извлечены по ключевым словам: {width}x{height}x{depth}")
            else:
                # Извлекаем все числа из текста размеров (старый метод)
                all_numbers = re.findall(r'\d+', size_text)
                print(f"   📊 Найдено чисел: {all_numbers}")
                if len(all_numbers) >= 2:
                    width = int(all_numbers[0])
                    height = int(all_numbers[1])
                    depth = int(all_numbers[2]) if len(all_numbers) > 2 else 600
                    print(f"   ✅ Размеры извлечены из списка чисел: {width}x{height}x{depth}")
                else:
                    print(f"   ⚠️ Недостаточно чисел в размерах, используем значения по умолчанию")
        except Exception as e:
            print(f"   ⚠️ Ошибка парсинга размеров: {e}")
            print(f"   Используем значения по умолчанию: 0x0x600")
        
        # Извлекаем материалы
        body_text = cells[body_idx].text.strip().lower() if body_idx < len(cells) else ""
        facade_text = cells[facade_idx].text.strip().lower() if facade_idx < len(cells) else ""
        
        item_data = {
            "name": item_name,
            "width": width,
            "height": height,
            "depth": depth,
            "body": body_text,
            "facade": facade_text,
            "image": image_path,
            "quantity": quantity
        }
        print(f"   📐 Размеры: {item_data['width']}x{item_data['height']}x{item_data['depth']} мм")
        print(f"   📦 Корпус: {item_data['body']} | Фасад: {item_data['facade']}")
        extracted_items.append(item_data)
            
    return extracted_items, furniture_items

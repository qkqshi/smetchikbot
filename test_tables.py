import docx

doc = docx.Document('uploads/пример кухни (2).docx')
print(f'Таблиц в документе: {len(doc.tables)}')

for i, table in enumerate(doc.tables):
    print(f'\n=== Таблица {i+1} ===')
    print(f'Строк: {len(table.rows)}, Столбцов: {len(table.rows[0].cells)}')
    print('Заголовки:')
    for cell in table.rows[0].cells:
        print(f'  - {cell.text}')
    
    if len(table.rows) > 1:
        print('Первая строка данных:')
        for cell in table.rows[1].cells:
            print(f'  - {cell.text[:100]}')

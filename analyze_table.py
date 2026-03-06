import docx

doc = docx.Document('uploads/ВЯЗНИКОВСКАЯ просчет.docx')
print(f'Таблиц в документе: {len(doc.tables)}')

if doc.tables:
    table = doc.tables[0]
    print(f'Строк в таблице: {len(table.rows)}')
    print(f'Столбцов: {len(table.rows[0].cells)}')
    
    print('\n=== ЗАГОЛОВКИ ===')
    for i, cell in enumerate(table.rows[0].cells):
        print(f'Столбец {i}: "{cell.text.strip()}"')
    
    print('\n=== ПЕРВАЯ СТРОКА ДАННЫХ ===')
    if len(table.rows) > 1:
        for i, cell in enumerate(table.rows[1].cells):
            text = cell.text.strip()
            if len(text) > 100:
                text = text[:100] + '...'
            print(f'Столбец {i}: "{text}"')
    
    print('\n=== ВТОРАЯ СТРОКА ДАННЫХ ===')
    if len(table.rows) > 2:
        for i, cell in enumerate(table.rows[2].cells):
            text = cell.text.strip()
            if len(text) > 100:
                text = text[:100] + '...'
            print(f'Столбец {i}: "{text}"')
